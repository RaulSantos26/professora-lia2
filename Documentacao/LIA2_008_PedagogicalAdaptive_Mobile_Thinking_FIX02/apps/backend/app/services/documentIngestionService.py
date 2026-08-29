import io
import zipfile
from pathlib import Path
from uuid import UUID

import pypdfium2 as pdfium
from PIL import Image, ImageOps
from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.persistence.models.documentBlockModel import DocumentBlockModel
from app.persistence.models.documentChunkModel import DocumentChunkModel
from app.persistence.models.documentModel import DocumentModel
from app.persistence.models.documentPageModel import DocumentPageModel
from app.persistence.models.documentVersionModel import DocumentVersionModel
from app.persistence.models.evidenceModel import EvidenceModel
from app.repositories.documentRepository import DocumentRepository
from app.services.materialStorageService import MaterialStorageService


class DocumentProcessingError(Exception):
    def __init__(self, code: str, safeMessage: str):
        super().__init__(safeMessage)
        self.code = code
        self.safeMessage = safeMessage


class DocumentIngestionService:
    CHUNK_SIZE = 1500

    def __init__(self, session: Session):
        self.session = session
        self.repository = DocumentRepository(session)
        self.storage = MaterialStorageService()

    def ingest(
        self,
        studentId: UUID,
        materialId: UUID,
        materialFileId: UUID,
        materialType: str,
        filePath: Path,
    ) -> tuple[
        DocumentModel,
        DocumentVersionModel,
        int,
        int,
        int,
    ]:
        document = DocumentModel(
            materialId=materialId,
            status="PROCESSING",
            pageCount=0,
        )
        self.repository.createDocument(document)

        version = DocumentVersionModel(
            documentId=document.documentId,
            materialFileId=materialFileId,
            versionNumber=1,
            extractionStatus="PENDING",
        )
        self.repository.createVersion(version)

        if materialType == "PDF":
            return self._ingestPdf(
                studentId,
                materialId,
                document,
                version,
                filePath,
            )

        if materialType == "IMAGE":
            return self._ingestImage(
                studentId,
                materialId,
                document,
                version,
                filePath,
            )

        if materialType == "TEXT":
            return self._ingestTextFile(
                studentId,
                materialId,
                document,
                version,
                filePath,
            )

        if materialType == "DOCUMENT":
            if filePath.suffix.lower() == ".docx":
                return self._ingestDocx(
                    studentId,
                    materialId,
                    document,
                    version,
                    filePath,
                )

            return self._markStructurePending(
                studentId,
                materialId,
                document,
                version,
                locator="documento 1",
            )

        return self._markStructurePending(
            studentId,
            materialId,
            document,
            version,
            locator="arquivo 1",
        )

    def _ingestPdf(
        self,
        studentId: UUID,
        materialId: UUID,
        document: DocumentModel,
        version: DocumentVersionModel,
        filePath: Path,
    ):
        try:
            reader = PdfReader(str(filePath))

            if reader.is_encrypted:
                try:
                    unlocked = reader.decrypt("")
                except Exception:
                    unlocked = 0

                if not unlocked:
                    raise DocumentProcessingError(
                        code="PDF_PASSWORD_PROTECTED",
                        safeMessage=(
                            "O PDF está protegido por senha e não pode "
                            "ser analisado automaticamente."
                        ),
                    )

            renderer = pdfium.PdfDocument(str(filePath))

            if len(renderer) != len(reader.pages):
                raise DocumentProcessingError(
                    code="PDF_PAGE_COUNT_MISMATCH",
                    safeMessage=(
                        "Os leitores de PDF retornaram quantidades "
                        "diferentes de páginas."
                    ),
                )

            document.pageCount = len(reader.pages)

            chunkIndex = 0
            textBlocks = 0
            visualPending = 0

            for pageIndex, pdfPage in enumerate(
                reader.pages,
                start=1,
            ):
                try:
                    nativeText = (
                        pdfPage.extract_text() or ""
                    ).strip()
                except Exception:
                    nativeText = ""

                renderPage = renderer[pageIndex - 1]

                # Correctness-first policy for LIA2-007:
                # every PDF page is also preserved as a visual source.
                # Native text and visual interpretation are independent
                # evidence channels; this prevents diagrams/vector figures
                # from disappearing merely because the page has text.
                requiresVision = True

                page = DocumentPageModel(
                    documentVersionId=version.documentVersionId,
                    pageNumber=pageIndex,
                    nativeText=nativeText or None,
                    status=(
                        "VISUAL_PENDING"
                        if requiresVision
                        else (
                            "TEXT_READY"
                            if nativeText
                            else "EMPTY"
                        )
                    ),
                )
                self.repository.createPage(page)

                sequence = 1

                if nativeText:
                    block, evidence = self._createTextBlock(
                        studentId=studentId,
                        materialId=materialId,
                        version=version,
                        page=page,
                        sequenceNumber=sequence,
                        text=nativeText,
                        locator=f"p. {pageIndex} · texto nativo",
                    )
                    textBlocks += 1
                    sequence += 1

                    for chunk in self._chunks(nativeText):
                        self.repository.createChunk(
                            DocumentChunkModel(
                                documentVersionId=version.documentVersionId,
                                documentPageId=page.documentPageId,
                                documentBlockId=block.documentBlockId,
                                evidenceId=evidence.evidenceId,
                                chunkIndex=chunkIndex,
                                content=chunk,
                                tokenEstimate=max(
                                    1,
                                    len(chunk) // 4,
                                ),
                                status="PENDING_EMBEDDING",
                            )
                        )
                        chunkIndex += 1

                if requiresVision:
                    bitmap = renderPage.render(scale=1.5)
                    pilImage = bitmap.to_pil()
                    output = io.BytesIO()
                    pilImage.convert("RGB").save(
                        output,
                        format="PNG",
                        optimize=True,
                    )

                    assetKey = self.storage.saveDerivedBytes(
                        studentId=studentId,
                        materialId=materialId,
                        relativeName=(
                            f"pdf/page-{pageIndex:04d}.png"
                        ),
                        content=output.getvalue(),
                    )

                    self._createVisualBlock(
                        studentId=studentId,
                        materialId=materialId,
                        version=version,
                        page=page,
                        sequenceNumber=sequence,
                        assetStorageKey=assetKey,
                        locator=f"p. {pageIndex} · render visual",
                    )
                    visualPending += 1

            renderer.close()

            if visualPending:
                document.status = "PARTIAL"
                version.extractionStatus = (
                    "PARTIAL"
                    if textBlocks > 0
                    else "VISUAL_PENDING"
                )
            elif textBlocks:
                document.status = "READY"
                version.extractionStatus = "NATIVE_TEXT_READY"
            else:
                document.status = "PARTIAL"
                version.extractionStatus = "PENDING"

            return (
                document,
                version,
                textBlocks,
                visualPending,
                chunkIndex,
            )

        except DocumentProcessingError:
            raise

        except Exception as error:
            raise DocumentProcessingError(
                code="PDF_PARSE_ERROR",
                safeMessage=(
                    "Não foi possível interpretar este PDF. "
                    "O arquivo original foi preservado."
                ),
            ) from error

    def _ingestImage(
        self,
        studentId: UUID,
        materialId: UUID,
        document: DocumentModel,
        version: DocumentVersionModel,
        filePath: Path,
    ):
        try:
            with Image.open(filePath) as source:
                normalized = ImageOps.exif_transpose(source)
                output = io.BytesIO()
                normalized.convert("RGB").save(
                    output,
                    format="PNG",
                    optimize=True,
                )

            assetKey = self.storage.saveDerivedBytes(
                studentId=studentId,
                materialId=materialId,
                relativeName="image/source-normalized.png",
                content=output.getvalue(),
            )

            document.pageCount = 1

            page = DocumentPageModel(
                documentVersionId=version.documentVersionId,
                pageNumber=1,
                nativeText=None,
                status="VISUAL_PENDING",
            )
            self.repository.createPage(page)

            self._createVisualBlock(
                studentId=studentId,
                materialId=materialId,
                version=version,
                page=page,
                sequenceNumber=1,
                assetStorageKey=assetKey,
                locator="imagem 1",
            )

            document.status = "PARTIAL"
            version.extractionStatus = "VISUAL_PENDING"

            return document, version, 0, 1, 0

        except Exception as error:
            raise DocumentProcessingError(
                code="IMAGE_PREPARE_ERROR",
                safeMessage=(
                    "Não foi possível preparar a imagem para análise."
                ),
            ) from error

    def _ingestDocx(
        self,
        studentId: UUID,
        materialId: UUID,
        document: DocumentModel,
        version: DocumentVersionModel,
        filePath: Path,
    ):
        try:
            docx = DocxDocument(str(filePath))

            textParts = []

            for paragraph in docx.paragraphs:
                value = paragraph.text.strip()

                if value:
                    textParts.append(value)

            for table in docx.tables:
                for row in table.rows:
                    values = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]

                    if values:
                        textParts.append(" | ".join(values))

            nativeText = "\n".join(textParts).strip()
            mediaItems = self._extractDocxMedia(filePath)

            document.pageCount = 1

            page = DocumentPageModel(
                documentVersionId=version.documentVersionId,
                pageNumber=1,
                nativeText=nativeText or None,
                status=(
                    "VISUAL_PENDING"
                    if mediaItems
                    else (
                        "TEXT_READY"
                        if nativeText
                        else "EMPTY"
                    )
                ),
            )
            self.repository.createPage(page)

            sequence = 1
            chunkIndex = 0
            textBlocks = 0

            if nativeText:
                block, evidence = self._createTextBlock(
                    studentId=studentId,
                    materialId=materialId,
                    version=version,
                    page=page,
                    sequenceNumber=sequence,
                    text=nativeText,
                    locator="DOCX · texto nativo",
                )
                sequence += 1
                textBlocks = 1

                for chunk in self._chunks(nativeText):
                    self.repository.createChunk(
                        DocumentChunkModel(
                            documentVersionId=version.documentVersionId,
                            documentPageId=page.documentPageId,
                            documentBlockId=block.documentBlockId,
                            evidenceId=evidence.evidenceId,
                            chunkIndex=chunkIndex,
                            content=chunk,
                            tokenEstimate=max(
                                1,
                                len(chunk) // 4,
                            ),
                            status="PENDING_EMBEDDING",
                        )
                    )
                    chunkIndex += 1

            visualPending = 0

            for mediaIndex, (mediaName, mediaBytes) in enumerate(
                mediaItems,
                start=1,
            ):
                suffix = Path(mediaName).suffix.lower()

                if suffix not in {
                    ".png",
                    ".jpg",
                    ".jpeg",
                    ".webp",
                    ".bmp",
                    ".gif",
                }:
                    continue

                with Image.open(io.BytesIO(mediaBytes)) as mediaImage:
                    normalizedMedia = ImageOps.exif_transpose(
                        mediaImage
                    )
                    mediaOutput = io.BytesIO()
                    normalizedMedia.convert("RGB").save(
                        mediaOutput,
                        format="PNG",
                        optimize=True,
                    )

                assetKey = self.storage.saveDerivedBytes(
                    studentId=studentId,
                    materialId=materialId,
                    relativeName=(
                        f"docx/image-{mediaIndex:04d}.png"
                    ),
                    content=mediaOutput.getvalue(),
                )

                self._createVisualBlock(
                    studentId=studentId,
                    materialId=materialId,
                    version=version,
                    page=page,
                    sequenceNumber=sequence,
                    assetStorageKey=assetKey,
                    locator=(
                        f"DOCX · imagem incorporada {mediaIndex}"
                    ),
                )
                sequence += 1
                visualPending += 1

            if visualPending:
                document.status = "PARTIAL"
                version.extractionStatus = (
                    "PARTIAL"
                    if textBlocks
                    else "VISUAL_PENDING"
                )
            elif textBlocks:
                document.status = "READY"
                version.extractionStatus = "NATIVE_TEXT_READY"
            else:
                document.status = "PARTIAL"
                version.extractionStatus = "PENDING"

            return (
                document,
                version,
                textBlocks,
                visualPending,
                chunkIndex,
            )

        except Exception as error:
            raise DocumentProcessingError(
                code="DOCX_PARSE_ERROR",
                safeMessage=(
                    "Não foi possível interpretar este DOCX. "
                    "O arquivo original foi preservado."
                ),
            ) from error

    def _extractDocxMedia(
        self,
        filePath: Path,
    ) -> list[tuple[str, bytes]]:
        items = []

        with zipfile.ZipFile(filePath, "r") as archive:
            for name in sorted(archive.namelist()):
                if not name.startswith("word/media/"):
                    continue

                items.append(
                    (
                        Path(name).name,
                        archive.read(name),
                    )
                )

        return items

    def _ingestTextFile(
        self,
        studentId: UUID,
        materialId: UUID,
        document: DocumentModel,
        version: DocumentVersionModel,
        filePath: Path,
    ):
        try:
            text = filePath.read_text(
                encoding="utf-8",
                errors="replace",
            ).strip()
        except Exception as error:
            raise DocumentProcessingError(
                code="TEXT_PARSE_ERROR",
                safeMessage=(
                    "Não foi possível ler este arquivo de texto."
                ),
            ) from error

        document.pageCount = 1

        page = DocumentPageModel(
            documentVersionId=version.documentVersionId,
            pageNumber=1,
            nativeText=text or None,
            status="TEXT_READY" if text else "EMPTY",
        )
        self.repository.createPage(page)

        if not text:
            document.status = "PARTIAL"
            version.extractionStatus = "PENDING"
            return document, version, 0, 0, 0

        block, evidence = self._createTextBlock(
            studentId=studentId,
            materialId=materialId,
            version=version,
            page=page,
            sequenceNumber=1,
            text=text,
            locator="texto 1",
        )

        chunkIndex = 0

        for chunk in self._chunks(text):
            self.repository.createChunk(
                DocumentChunkModel(
                    documentVersionId=version.documentVersionId,
                    documentPageId=page.documentPageId,
                    documentBlockId=block.documentBlockId,
                    evidenceId=evidence.evidenceId,
                    chunkIndex=chunkIndex,
                    content=chunk,
                    tokenEstimate=max(
                        1,
                        len(chunk) // 4,
                    ),
                    status="PENDING_EMBEDDING",
                )
            )
            chunkIndex += 1

        document.status = "READY"
        version.extractionStatus = "NATIVE_TEXT_READY"

        return document, version, 1, 0, chunkIndex

    def _markStructurePending(
        self,
        studentId: UUID,
        materialId: UUID,
        document: DocumentModel,
        version: DocumentVersionModel,
        locator: str,
    ):
        document.pageCount = 1

        page = DocumentPageModel(
            documentVersionId=version.documentVersionId,
            pageNumber=1,
            nativeText=None,
            status="EMPTY",
        )
        self.repository.createPage(page)

        block = DocumentBlockModel(
            documentPageId=page.documentPageId,
            sequenceNumber=1,
            blockType="OTHER",
            textContent=None,
            processingStatus="PENDING_STRUCTURE",
        )
        self.repository.createBlock(block)

        self.repository.createEvidence(
            EvidenceModel(
                studentId=studentId,
                materialId=materialId,
                documentVersionId=version.documentVersionId,
                documentPageId=page.documentPageId,
                documentBlockId=block.documentBlockId,
                evidenceType="DOCUMENT",
                locator=locator,
                excerpt=None,
                status="ACTIVE",
            )
        )

        document.status = "PARTIAL"
        version.extractionStatus = "PENDING"

        return document, version, 0, 0, 0

    def _createTextBlock(
        self,
        *,
        studentId: UUID,
        materialId: UUID,
        version: DocumentVersionModel,
        page: DocumentPageModel,
        sequenceNumber: int,
        text: str,
        locator: str,
    ):
        block = DocumentBlockModel(
            documentPageId=page.documentPageId,
            sequenceNumber=sequenceNumber,
            blockType="TEXT",
            textContent=text,
            processingStatus="READY",
        )
        self.repository.createBlock(block)

        evidence = EvidenceModel(
            studentId=studentId,
            materialId=materialId,
            documentVersionId=version.documentVersionId,
            documentPageId=page.documentPageId,
            documentBlockId=block.documentBlockId,
            evidenceType="TEXT",
            locator=locator,
            excerpt=text[:1000],
            status="ACTIVE",
        )
        self.repository.createEvidence(evidence)

        return block, evidence

    def _createVisualBlock(
        self,
        *,
        studentId: UUID,
        materialId: UUID,
        version: DocumentVersionModel,
        page: DocumentPageModel,
        sequenceNumber: int,
        assetStorageKey: str,
        locator: str,
    ) -> DocumentBlockModel:
        block = DocumentBlockModel(
            documentPageId=page.documentPageId,
            sequenceNumber=sequenceNumber,
            blockType="IMAGE",
            textContent=None,
            processingStatus="PENDING_VISION",
            assetStorageKey=assetStorageKey,
        )
        self.repository.createBlock(block)

        self.repository.createEvidence(
            EvidenceModel(
                studentId=studentId,
                materialId=materialId,
                documentVersionId=version.documentVersionId,
                documentPageId=page.documentPageId,
                documentBlockId=block.documentBlockId,
                evidenceType="IMAGE",
                locator=locator,
                excerpt=None,
                status="ACTIVE",
            )
        )

        return block

    def _chunks(self, text: str) -> list[str]:
        normalized = "\n".join(
            line.strip()
            for line in text.splitlines()
            if line.strip()
        )

        if not normalized:
            return []

        chunks = []
        cursor = 0

        while cursor < len(normalized):
            end = min(
                cursor + self.CHUNK_SIZE,
                len(normalized),
            )

            if end < len(normalized):
                split = normalized.rfind(
                    "\n",
                    cursor,
                    end,
                )

                if split > cursor + 300:
                    end = split

            chunk = normalized[cursor:end].strip()

            if chunk:
                chunks.append(chunk)

            cursor = max(end, cursor + 1)

        return chunks
